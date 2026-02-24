"""Convert MANUSCRIPT_DRAFT_v1.md to Word docx for co-author review."""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

MANUSCRIPT = Path('C:/Users/moshe/Dropbox/ISF 2025/agent_coordination/MANUSCRIPT_DRAFT_v4.md')
OUTPUT     = Path('C:/Users/moshe/Dropbox/ISF 2025/agent_coordination/MANUSCRIPT_DRAFT_v4.docx')
FIGS_DIR   = Path('C:/Users/moshe/Dropbox/ISF 2025/state_space_figures')

FIG_FILES = {
    '1':  'Fig1_framework_schematic.png',
    '2':  'paper_heatmap_5tx_10axes.png',
    '3':  'PGPR_split_JA_SA_4panel.png',
    '4':  'GSE87337_PHR1_validation.png',
    '5':  'Fig_reproducibility_amino_acids.png',
    '6':  'Fig_PSK_origin_classification_failure.png',
    '7':  'coherence_analysis.png',
    '8':  'cross_species_conservation_barplot.png',
    'S1': 'GSE271932_Pi_resupply_timecourse.png',
    'S2': 'axis_correlation_matrix.png',
    'S3': 'axis_Fstat_ranking.png',
    'S4': 'PCA_5treatments_9axes.png',
    'S5': 'discriminant_scatter_paper.png',
    'S6': 'classification_dendrogram.png',
    'S7': 'Fig_seaweed_prediction_heatmap.png',
    'S8': 'amino_acid_gene_level_reproducibility.png',
}

# ---------- read & trim ----------
with open(MANUSCRIPT, encoding='utf-8') as f:
    raw = f.readlines()

STOP_MARKERS = {'## WORD COUNT SUMMARY', '## ITEMS STILL NEEDED BEFORE SUBMISSION'}
SKIP_PREFIXES = ('*MANUSCRIPT_DRAFT', '*Revised', '*Source files', '*All source files')

lines = []
for ln in raw:
    s = ln.rstrip('\n')
    if s.strip() in STOP_MARKERS:
        break
    if any(s.strip().startswith(p) for p in SKIP_PREFIXES):
        continue
    lines.append(s)

# ---------- helpers ----------
def parse_inline(text):
    """Return list of (text, bold, italic, code)."""
    result = []
    pat = r'(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*\n]+\*|`[^`]+`)'
    for part in re.split(pat, text):
        if not part:
            continue
        if part.startswith('***') and part.endswith('***'):
            result.append((part[3:-3], True, True, False))
        elif part.startswith('**') and part.endswith('**'):
            result.append((part[2:-2], True, False, False))
        elif part.startswith('*') and part.endswith('*'):
            result.append((part[1:-1], False, True, False))
        elif part.startswith('`') and part.endswith('`'):
            result.append((part[1:-1], False, False, True))
        else:
            result.append((part, False, False, False))
    return result

def add_para(doc, text, style='Normal', heading_level=None,
             centered=False, bold_all=False, italic_all=False, gray=False):
    wstyle = f'Heading {heading_level}' if heading_level else style
    p = doc.add_paragraph(style=wstyle)
    if centered:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Replace -- with em dash
    text = text.replace(' -- ', '\u2014').replace('--', '\u2014')
    for t, b, it, c in parse_inline(text):
        if not t:
            continue
        run = p.add_run(t)
        run.bold  = b or bold_all
        run.italic = it or italic_all
        if c:
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
        if gray:
            run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    return p

def fill_cell(cell, text, header=False):
    p = cell.paragraphs[0]
    p.clear()
    text = text.replace(' -- ', '\u2014').replace('--', '\u2014')
    for t, b, it, c in parse_inline(text):
        if not t:
            continue
        run = p.add_run(t)
        run.bold   = b or header
        run.italic = it
        run.font.size = Pt(9)
        if c:
            run.font.name = 'Courier New'

def flush_para(doc, buf):
    if not buf:
        return
    text = ' '.join(b.strip() for b in buf if b.strip())
    if text:
        add_para(doc, text)
    buf.clear()

def flush_table(doc, rows):
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = 'Table Grid'
    for ri, row in enumerate(rows):
        for ci, ct in enumerate(row):
            if ci < ncols:
                fill_cell(tbl.cell(ri, ci), ct, header=(ri == 0))
    # spacing after
    doc.add_paragraph()

# ---------- build document ----------
doc = Document()
for sec in doc.sections:
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.25)
    sec.right_margin  = Inches(1.25)

ns = doc.styles['Normal']
ns.font.name = 'Times New Roman'
ns.font.size = Pt(12)
ns.paragraph_format.space_after = Pt(6)

para_buf   = []
table_rows = []
in_table   = False
in_eq      = False
eq_buf     = []
main_figs_inserted = False

i = 0
while i < len(lines):
    raw_line = lines[i]
    stripped = raw_line.strip()
    i += 1

    # --- equation block ---
    if stripped == '$$':
        flush_para(doc, para_buf)
        if in_eq:
            eq_text = ' '.join(eq_buf)
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run(eq_text).font.name = 'Courier New'
            in_eq = False; eq_buf = []
        else:
            in_eq = True
        continue

    if stripped.startswith('$$') and stripped.endswith('$$') and len(stripped) > 4:
        flush_para(doc, para_buf)
        eq_text = stripped[2:-2]
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(eq_text).font.name = 'Courier New'
        continue

    if in_eq:
        eq_buf.append(stripped)
        continue

    # --- table row ---
    if stripped.startswith('|'):
        flush_para(doc, para_buf)
        if re.match(r'^\|[\s\-:|]+\|', stripped):
            continue  # separator row
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        table_rows.append(cells)
        in_table = True
        continue

    # end of table
    if in_table:
        flush_table(doc, table_rows)
        table_rows = []
        in_table = False

    # --- blank line ---
    if not stripped:
        flush_para(doc, para_buf)
        continue

    # --- headings ---
    if stripped.startswith('#### '):
        flush_para(doc, para_buf)
        add_para(doc, stripped[5:], heading_level=4)
        continue
    if stripped.startswith('### '):
        flush_para(doc, para_buf)
        heading_text = stripped[4:]
        add_para(doc, heading_text, heading_level=3)
        # Embed supplementary figure images inline after their heading
        m = re.match(r'^Figure (S\d+)\.', heading_text)
        if m:
            fig_key = m.group(1)
            fname = FIG_FILES.get(fig_key)
            if fname:
                fig_path = FIGS_DIR / fname
                if fig_path.exists():
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(str(fig_path), width=Inches(5.5))
        continue
    if stripped.startswith('## '):
        flush_para(doc, para_buf)
        add_para(doc, stripped[3:], heading_level=2)
        continue
    if stripped.startswith('# '):
        flush_para(doc, para_buf)
        add_para(doc, stripped[2:], heading_level=1)
        continue

    # --- horizontal rule ---
    if re.match(r'^-{3,}$', stripped):
        flush_para(doc, para_buf)
        doc.add_paragraph()
        continue

    # --- figure / table placeholders ---
    if re.match(r'^\[(?:FIGURE|TABLE)', stripped):
        flush_para(doc, para_buf)
        m = re.match(r'^\[FIGURE\s*(S?\d+)\]', stripped)
        if m:
            fig_key = m.group(1)
            if not fig_key.startswith('S'):
                # Main figures: on first encounter (Figure 1), dump all 8 in order
                if not main_figs_inserted:
                    main_figs_inserted = True
                    add_para(doc, 'FIGURES', heading_level=2)
                    for fnum in ['1','2','3','4','5','6','7','8']:
                        fname = FIG_FILES.get(fnum)
                        if fname:
                            fig_path = FIGS_DIR / fname
                            if fig_path.exists():
                                p = doc.add_paragraph()
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                p.add_run().add_picture(str(fig_path), width=Inches(5.5))
                                cap = doc.add_paragraph(style='Normal')
                                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                r = cap.add_run(f'Figure {fnum}')
                                r.bold = True
                                r.font.size = Pt(10)
                                doc.add_paragraph()
                # Skip subsequent main figure placeholders
                continue
        add_para(doc, stripped, italic_all=True, gray=True)
        continue

    # --- bullet points ---
    if re.match(r'^-\s+\S', stripped) or re.match(r'^  -\s+\S', stripped):
        flush_para(doc, para_buf)
        bullet_text = re.sub(r'^-\s+', '', stripped.lstrip())
        add_para(doc, bullet_text, style='List Bullet')
        continue

    # --- everything else: buffer ---
    para_buf.append(stripped)

# final flush
flush_para(doc, para_buf)
if table_rows:
    flush_table(doc, table_rows)

doc.save(str(OUTPUT))
print(f'Saved: {OUTPUT}')
print(f'Size: {OUTPUT.stat().st_size / 1024:.0f} KB')
